#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate_review_docx.py v1.0 — Xuất PHIẾU ĐÁNH GIÁ BÁO CÁO TVGS ra file .docx A4
(Phòng Kỹ thuật — Công ty Cổ phần TEXO Tư vấn và Đầu tư)

Input: file JSON do Claude tổng hợp từ kết quả Tầng 1 + Tầng 2 (xem review_data_template.json)
Usage:
    python generate_review_docx.py review_data.json -o "Danh gia bao cao ....docx"

Dependencies: python-docx
"""

import sys, json, argparse, datetime
from docx import Document
from docx.shared import Pt, Mm, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----- bảng màu (hiện đại, nhận diện TEXO: xanh đậm + cam) -----
C_PRIMARY   = '1F4E79'   # xanh đậm
C_ACCENT    = 'E8731A'   # cam
C_GREY      = '595959'
C_LIGHT     = 'EAF1F8'   # nền xanh nhạt
C_OK        = '2E7D32'; C_OK_BG   = 'E2F0E4'
C_WARN      = 'B7791F'; C_WARN_BG = 'FDF3D7'
C_BAD       = 'C0392B'; C_BAD_BG  = 'FADBD8'

STATUS = {
    'DAT':     ('ĐẠT', C_OK, C_OK_BG),
    'CAN_SUA': ('CẦN SỬA', C_WARN, C_WARN_BG),
    'LOI':     ('LỖI', C_BAD, C_BAD_BG),
    'THIEU':   ('THIẾU MỤC', C_BAD, C_BAD_BG),
}
XEP_LOAI_COLOR = {
    'ĐẠT': (C_OK, C_OK_BG), 'ĐẠT CÓ ĐIỀU KIỆN': (C_WARN, C_WARN_BG),
    'CẦN SỬA': (C_WARN, C_WARN_BG), 'CẦN SỬA NHIỀU': (C_WARN, C_WARN_BG),
    'KHÔNG ĐẠT': (C_BAD, C_BAD_BG),
}

FONT = 'Calibri'


def set_font(run, size=11, bold=False, color=None, italic=False, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    ea = rpr.find(qn('w:rFonts'))
    if ea is None:
        ea = OxmlElement('w:rFonts'); rpr.append(ea)
    ea.set(qn('w:ascii'), font); ea.set(qn('w:hAnsi'), font); ea.set(qn('w:cs'), font)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def cell_margins(table, top=60, bottom=60, left=110, right=110):
    tblPr = table._tbl.tblPr
    m = OxmlElement('w:tblCellMar')
    for tag, v in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        e = OxmlElement(f'w:{tag}'); e.set(qn('w:w'), str(v)); e.set(qn('w:type'), 'dxa')
        m.append(e)
    tblPr.append(m)


def table_borders(table, color='BFD3E6', size='4'):
    tblPr = table._tbl.tblPr
    b = OxmlElement('w:tblBorders')
    for tag in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{tag}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), size)
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        b.append(e)
    tblPr.append(b)


def no_borders(table):
    tblPr = table._tbl.tblPr
    b = OxmlElement('w:tblBorders')
    for tag in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{tag}'); e.set(qn('w:val'), 'none')
        b.append(e)
    tblPr.append(b)


def para_border_bottom(p, color=C_PRIMARY, size='12'):
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), size)
    bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), color)
    pbdr.append(bot); ppr.append(pbdr)


def para_border_top(p, color='BFD3E6', size='6'):
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single'); top.set(qn('w:sz'), size)
    top.set(qn('w:space'), '4'); top.set(qn('w:color'), color)
    pbdr.append(top); ppr.append(pbdr)


def right_tab(p, cm=16.8):
    """Vô hiệu tab mặc định của style Header/Footer (4513/9026 twips A4, 4680/9360 Letter)
    bằng w:tab val='clear', rồi đặt 1 tab phải duy nhất."""
    ppr = p._p.get_or_add_pPr()
    old = ppr.find(qn('w:tabs'))
    if old is not None:
        ppr.remove(old)
    tabs = OxmlElement('w:tabs')
    for pos in ('4513', '9026', '4680', '9360'):
        e = OxmlElement('w:tab')
        e.set(qn('w:val'), 'clear'); e.set(qn('w:pos'), pos)
        tabs.append(e)
    e = OxmlElement('w:tab')
    e.set(qn('w:val'), 'right'); e.set(qn('w:pos'), str(int(cm * 567)))
    tabs.append(e)
    ppr.append(tabs)


def row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    e = OxmlElement('w:cantSplit'); trPr.append(e)


def add_field(p, instr):
    """Chèn field code (PAGE / NUMPAGES) vào paragraph."""
    r = p.add_run()
    for el, attrs in (('w:fldChar', {'w:fldCharType': 'begin'}),
                      ('w:instrText', None),
                      ('w:fldChar', {'w:fldCharType': 'end'})):
        e = OxmlElement(el)
        if attrs:
            for k, v in attrs.items():
                e.set(qn(k), v)
        if el == 'w:instrText':
            e.set(qn('xml:space'), 'preserve'); e.text = f' {instr} '
        r._element.append(e)
    set_font(r, 9, color=C_GREY)


def spacing(p, before=0, after=6, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    if line:
        pf.line_spacing = line


def heading(doc, num, text):
    p = doc.add_paragraph()
    r1 = p.add_run(f'{num}  ')
    set_font(r1, 14, bold=True, color=C_ACCENT)
    r2 = p.add_run(text.upper())
    set_font(r2, 14, bold=True, color=C_PRIMARY)
    para_border_bottom(p, color='BFD3E6', size='6')
    spacing(p, before=14, after=8)
    return p


def bullet(doc, text, color=None, size=10.5, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    spacing(p, after=3)
    r0 = p.add_run('— ' if color is None else '■ ')
    set_font(r0, size, color=color or C_PRIMARY, bold=True)
    if bold_prefix:
        rb = p.add_run(bold_prefix + ' ')
        set_font(rb, size, bold=True, color=color or '262626')
    r = p.add_run(text)
    set_font(r, size, color='262626' if color is None else color)
    return p


def generate(d, output):
    """Render review_data dict -> file .docx (dùng được cả từ app lẫn CLI)."""

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin = Cm(2.6); sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(1.0); sec.footer_distance = Cm(1.0)

    # ---------- HEADER ----------
    hp = sec.header.paragraphs[0]
    right_tab(hp)
    r = hp.add_run('CÔNG TY CỔ PHẦN TEXO TƯ VẤN VÀ ĐẦU TƯ')
    set_font(r, 10, bold=True, color=C_PRIMARY)
    r = hp.add_run('\tPHIẾU ĐÁNH GIÁ BÁO CÁO TVGS')
    set_font(r, 9, bold=True, color=C_ACCENT)
    h2 = sec.header.add_paragraph()
    right_tab(h2)
    r = h2.add_run('PHÒNG KỸ THUẬT')
    set_font(r, 9, bold=True, color=C_GREY)
    r = h2.add_run('\t' + d.get('so_phieu', ''))
    set_font(r, 9, color=C_GREY)
    para_border_bottom(h2, color=C_PRIMARY, size='14')
    spacing(hp, after=0); spacing(h2, after=0)

    # ---------- FOOTER ----------
    fp = sec.footer.paragraphs[0]
    right_tab(fp)
    para_border_top(fp, color=C_PRIMARY, size='8')
    r = fp.add_run(f"Phòng Kỹ thuật — TEXO  •  {d.get('ngay_danh_gia', '')}")
    set_font(r, 9, color=C_GREY)
    r = fp.add_run('\tTrang ')
    set_font(r, 9, color=C_GREY)
    add_field(fp, 'PAGE')
    r = fp.add_run('/')
    set_font(r, 9, color=C_GREY)
    add_field(fp, 'NUMPAGES')

    # ---------- KHỐI TIÊU ĐỀ ----------
    p = doc.add_paragraph(); spacing(p, before=0, after=2)
    r = p.add_run('PHÒNG KỸ THUẬT  •  ĐÁNH GIÁ CHẤT LƯỢNG BÁO CÁO TƯ VẤN GIÁM SÁT')
    set_font(r, 10, bold=True, color=C_ACCENT)
    p = doc.add_paragraph(); spacing(p, after=4)
    r = p.add_run(d.get('tieu_de', 'ĐÁNH GIÁ BÁO CÁO TVGS'))
    set_font(r, 19, bold=True, color=C_PRIMARY)
    para_border_bottom(p, color=C_ACCENT, size='18')

    # ---------- BẢNG THÔNG TIN ----------
    rows = [(k, v) for k, v in [
        ('Dự án / Công trình', d.get('du_an', '')),
        ('Hạng mục / Giai đoạn', d.get('hang_muc_giai_doan', '')),
        ('Loại báo cáo', d.get('loai_bao_cao_text', '')),
        ('Kỳ báo cáo', d.get('ky_bao_cao', '')),
        ('Báo cáo được đánh giá', d.get('file_bao_cao', '')),
        ('Đơn vị lập báo cáo', d.get('don_vi_lap', '')),
        ('Chủ đầu tư', d.get('chu_dau_tu', '')),
        ('Ngày đánh giá / Người đánh giá', f"{d.get('ngay_danh_gia','')} — {d.get('nguoi_danh_gia','Phòng Kỹ thuật')}"),
    ] if v]
    t = doc.add_table(rows=len(rows), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_borders(t, color='D9E4F0'); cell_margins(t)
    for i, (k, v) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        c0.width, c1.width = Cm(5.2), Cm(11.8)
        shade(c0, C_LIGHT)
        pr = c0.paragraphs[0]; r = pr.add_run(k); set_font(r, 10, bold=True, color=C_PRIMARY); spacing(pr, after=0)
        pr = c1.paragraphs[0]; r = pr.add_run(str(v)); set_font(r, 10); spacing(pr, after=0)

    # ---------- BANNER KẾT QUẢ ----------
    kq = d.get('ket_qua_chung', {})
    xl = kq.get('xep_loai', 'CẦN SỬA')
    fg, bg = XEP_LOAI_COLOR.get(xl.upper(), (C_WARN, C_WARN_BG))
    doc.add_paragraph()
    bt = doc.add_table(rows=1, cols=2)
    bt.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_borders(bt, color=fg, size='8'); cell_margins(bt, top=110, bottom=110)
    row_cant_split(bt.rows[0])
    c0, c1 = bt.rows[0].cells
    c0.width, c1.width = Cm(4.6), Cm(12.4)
    shade(c0, fg); shade(c1, bg)
    pr = c0.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.CENTER; spacing(pr, after=0)
    r = pr.add_run('KẾT QUẢ\n'); set_font(r, 9, bold=True, color='FFFFFF')
    r = pr.add_run(xl.upper()); set_font(r, 15, bold=True, color='FFFFFF')
    pr = c1.paragraphs[0]; spacing(pr, after=0)
    r = pr.add_run(kq.get('tom_tat', '')); set_font(r, 10.5, color='262626')

    # thống kê nhanh
    tk = d.get('thong_ke', {})
    if tk:
        p = doc.add_paragraph(); spacing(p, before=6, after=2)
        items = [(str(tk.get('dat', 0)), 'tiêu chí ĐẠT', C_OK),
                 (str(tk.get('can_sua', 0)), 'điểm CẦN SỬA', C_WARN),
                 (str(tk.get('loi', 0)), 'LỖI / mâu thuẫn', C_BAD),
                 (str(tk.get('muc_thieu', 0)), 'mục THIẾU', C_BAD)]
        for i, (n, lbl, cl) in enumerate(items):
            r = p.add_run(('     ' if i else '') + n + ' ')
            set_font(r, 13, bold=True, color=cl)
            r = p.add_run(lbl)
            set_font(r, 9.5, color=C_GREY)

    # ---------- 1. TỔNG HỢP THEO MỤC ----------
    heading(doc, '1', 'Tổng hợp đánh giá theo mục')
    dgm = d.get('danh_gia_muc', [])
    t = doc.add_table(rows=1 + len(dgm), cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_borders(t); cell_margins(t)
    widths = [Cm(1.5), Cm(5.4), Cm(2.5), Cm(7.6)]
    heads = ['Mục', 'Nội dung', 'Kết quả', 'Nhận xét chính']
    for j, htxt in enumerate(heads):
        c = t.rows[0].cells[j]; c.width = widths[j]
        shade(c, C_PRIMARY)
        pr = c.paragraphs[0]; spacing(pr, after=0)
        if j in (0, 2): pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pr.add_run(htxt); set_font(r, 10, bold=True, color='FFFFFF')
    for i, m in enumerate(dgm):
        cells = t.rows[i + 1].cells
        st_txt, st_fg, st_bg = STATUS.get(m.get('trang_thai', 'CAN_SUA'), STATUS['CAN_SUA'])
        vals = [m.get('muc', ''), m.get('ten', ''), st_txt, m.get('nhan_xet_ngan', '')]
        for j, v in enumerate(vals):
            c = cells[j]; c.width = widths[j]
            pr = c.paragraphs[0]; spacing(pr, after=0)
            if j in (0, 2): pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = pr.add_run(str(v))
            set_font(r, 9.5, bold=(j in (0, 2)),
                     color=st_fg if j == 2 else ('262626' if j != 0 else C_PRIMARY))
            if j == 2: shade(c, st_bg)
            if i % 2 == 1 and j != 2: shade(c, 'F5F8FC')

    # ---------- 2. PHÁT HIỆN CHÍNH ----------
    if d.get('phat_hien_chinh'):
        heading(doc, '2', 'Các phát hiện chính')
        for ph in d['phat_hien_chinh']:
            sev = ph.get('muc_do', 'CAN_SUA') if isinstance(ph, dict) else 'CAN_SUA'
            txt = ph.get('noi_dung', '') if isinstance(ph, dict) else str(ph)
            _, fg2, _ = STATUS.get(sev, STATUS['CAN_SUA'])
            bullet(doc, txt, color=fg2,
                   bold_prefix=f"[{STATUS.get(sev, STATUS['CAN_SUA'])[0]}]")

    # ---------- 3. ĐÁNH GIÁ CHI TIẾT ----------
    heading(doc, '3', 'Đánh giá chi tiết từng mục')
    for m in dgm:
        st_txt, st_fg, st_bg = STATUS.get(m.get('trang_thai', 'CAN_SUA'), STATUS['CAN_SUA'])
        p = doc.add_paragraph(); spacing(p, before=8, after=3)
        r = p.add_run(f"{m.get('muc','')} — {m.get('ten','')}   ")
        set_font(r, 11.5, bold=True, color=C_PRIMARY)
        r = p.add_run(f' {st_txt} ')
        set_font(r, 9.5, bold=True, color=st_fg)
        rpr = r._element.get_or_add_rPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), st_bg)
        rpr.append(shd)
        for nx in m.get('dat_duoc', []):
            bullet(doc, nx)
        for nx in m.get('van_de', []):
            bullet(doc, nx, color=st_fg if st_fg != C_OK else C_WARN)

    # ---------- 4. KIẾN NGHỊ ----------
    if d.get('kien_nghi'):
        heading(doc, '4', 'Yêu cầu chỉnh sửa & kiến nghị')
        for i, kn in enumerate(d['kien_nghi'], 1):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6); spacing(p, after=4)
            r = p.add_run(f'{i}. '); set_font(r, 10.5, bold=True, color=C_ACCENT)
            r = p.add_run(kn); set_font(r, 10.5)

    # ---------- 5. KẾT LUẬN ----------
    heading(doc, '5', 'Kết luận')
    kt = doc.add_table(rows=1, cols=1)
    table_borders(kt, color=C_PRIMARY, size='8'); cell_margins(kt, top=110, bottom=110)
    row_cant_split(kt.rows[0])
    c = kt.rows[0].cells[0]; c.width = Cm(17.0); shade(c, C_LIGHT)
    pr = c.paragraphs[0]; spacing(pr, after=0)
    r = pr.add_run(d.get('ket_luan', '')); set_font(r, 10.5, color='262626')

    # ---------- CHỮ KÝ ----------
    doc.add_paragraph()
    st = doc.add_table(rows=1, cols=2); no_borders(st)
    c0, c1 = st.rows[0].cells
    c0.width = c1.width = Cm(8.5)
    pr = c0.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.CENTER; spacing(pr, after=0)
    r = pr.add_run('NGƯỜI ĐÁNH GIÁ\n'); set_font(r, 10.5, bold=True, color=C_PRIMARY)
    r = pr.add_run('(Ký, ghi rõ họ tên)\n\n\n\n'); set_font(r, 9, italic=True, color=C_GREY)
    r = pr.add_run(d.get('nguoi_danh_gia_ten', '')); set_font(r, 10.5, bold=True)
    pr = c1.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.CENTER; spacing(pr, after=0)
    r = pr.add_run('TRƯỞNG PHÒNG KỸ THUẬT\n'); set_font(r, 10.5, bold=True, color=C_PRIMARY)
    r = pr.add_run('(Ký, ghi rõ họ tên)\n\n\n\n'); set_font(r, 9, italic=True, color=C_GREY)
    r = pr.add_run(d.get('truong_phong_ten', '')); set_font(r, 10.5, bold=True)

    doc.save(output)
    return output


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('review_json')
    ap.add_argument('-o', '--output', default='danh_gia_bao_cao_tvgs.docx')
    args = ap.parse_args()
    d = json.load(open(args.review_json, encoding='utf-8'))
    print('Đã xuất:', generate(d, args.output))


if __name__ == '__main__':
    main()
